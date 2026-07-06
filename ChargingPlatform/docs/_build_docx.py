"""One-shot: convert CZERO_OCPI_API_v1.0.md to a clean branded .docx.

Not a runtime dependency — kept in /docs as the source-of-truth generator
so the Word output can be regenerated whenever the markdown spec changes.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
SRC = ROOT / "CZERO_OCPI_API_v1.0.md"
OUT = ROOT / "CZERO_OCPI_API_v1.0.docx"

# Brand
BRAND_BLUE = RGBColor(0x0B, 0x4E, 0xA2)
BRAND_DARK = RGBColor(0x1A, 0x1A, 0x1A)
TABLE_HEAD = "0B4EA2"
TABLE_ALT  = "F2F5FA"
CODE_BG    = "F4F4F4"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level):
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = BRAND_BLUE if level <= 2 else BRAND_DARK
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)


INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")


def add_rich_paragraph(doc, text, style=None):
    p = doc.add_paragraph()
    if style:
        p.style = style
    # Resolve inline markdown: **bold** and `code`
    i = 0
    tokens = []
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            tokens.append(("plain", text[last:m.start()]))
        chunk = m.group(0)
        if chunk.startswith("**"):
            tokens.append(("bold", chunk[2:-2]))
        else:
            tokens.append(("code", chunk[1:-1]))
        last = m.end()
    if last < len(text):
        tokens.append(("plain", text[last:]))
    if not tokens:
        tokens = [("plain", text)]
    for kind, t in tokens:
        run = p.add_run(t)
        if kind == "bold":
            run.bold = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # Background shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)


def add_table(doc, rows):
    # Normalise ragged markdown tables — some rows may accidentally have
    # more or fewer pipes than the header. Take the max column count and
    # pad shorter rows with empty strings so cell access never overflows.
    max_cols = max(len(r) for r in rows) if rows else 0
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""  # clear default
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9.5)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, TABLE_HEAD)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, TABLE_ALT)
    return table


def parse_table(lines, start):
    """Parse a Markdown table starting at lines[start]. Return (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        # skip separator row of |---|---|
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        # strip inline backticks/bold for cleanliness (docx run doesn't need markdown)
        cells = [INLINE_CODE.sub(r"\1", BOLD.sub(r"\1", c)) for c in cells]
        rows.append(cells)
        i += 1
    return rows, i


def build(src_path=None, out_path=None, title=None, subtitle=None, date_str=None):
    """Convert a markdown file to a branded .docx.

    Kept backwards-compatible: called with no args, still builds the OCPI
    spec doc. Pass overrides to reuse for other partner-facing documents.
    """
    src = Path(src_path) if src_path else SRC
    out = Path(out_path) if out_path else OUT
    cover_title = title or "OCPI 2.2.1 API Specification"
    cover_sub = subtitle or "Partner Integration Reference — PlagSini EV Charging Platform"
    cover_date = date_str or "25 June 2026"

    md = src.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Page setup: A4 portrait, 2cm margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ── Cover ────────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(80)
    r = cover.add_run("C Zero Sdn Bhd")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = BRAND_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(cover_title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BRAND_DARK

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run(cover_sub)
    r.italic = True
    r.font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(60)
    for line in ("Version 1.0", cover_date, "Confidential — Partner Use Only"):
        r = meta.add_run(line + "\n")
        r.font.size = Pt(10)
    doc.add_page_break()

    # ── Body ─────────────────────────────────────────────────────────────
    i = 0
    in_code = False
    code_buffer = []
    while i < len(md):
        line = md[i]
        stripped = line.rstrip()

        # code fence
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), "0B4EA2")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # headings
        if stripped.startswith("# "):
            # skip the markdown H1 (we have cover); promote subsequent H1s
            if i == 0:
                i += 1; continue
            add_heading(doc, stripped[2:], 1); i += 1; continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1); i += 1; continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2); i += 1; continue
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:], 3); i += 1; continue

        # table
        if stripped.startswith("|"):
            rows, i = parse_table(md, i)
            if rows:
                add_table(doc, rows)
                doc.add_paragraph()  # spacing after table
            continue

        # bullet
        if re.match(r"^\s*[-*]\s+", stripped):
            text = re.sub(r"^\s*[-*]\s+", "", stripped)
            p = add_rich_paragraph(doc, text, style="List Bullet")
            i += 1; continue

        # ordered list
        if re.match(r"^\s*\d+\.\s+", stripped):
            text = re.sub(r"^\s*\d+\.\s+", "", stripped)
            p = add_rich_paragraph(doc, text, style="List Number")
            i += 1; continue

        # blank line
        if not stripped:
            i += 1; continue

        # plain paragraph
        # italic at start (* text *) → italic
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip("*"))
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1; continue
        add_rich_paragraph(doc, stripped)
        i += 1

    # Footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("C Zero Sdn Bhd · OCPI 2.2.1 Specification v1.0 · Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    # page number field
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run2 = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16"); rPr.append(sz)
    run2.append(rPr)
    t = OxmlElement("w:t"); t.text = "1"; run2.append(t)
    fld.append(run2)
    fp._p.append(fld)

    doc.save(str(out))
    print(f"Wrote {out}  ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "charger-req":
        build(
            src_path=ROOT / "CZero_22kW_Charger_Requirements_v1.0.md",
            out_path=ROOT / "CZero_22kW_Charger_Requirements_v1.0.docx",
            title="22kW AC Public EV Charger — Requirements",
            subtitle="Vendor Procurement Specification — C Zero Sdn Bhd",
            date_str="2 July 2026",
        )
    elif len(sys.argv) > 1 and sys.argv[1] == "transsemi":
        # Build the two partner-facing docs for Transsemi hardware onboarding.
        build(
            src_path=ROOT / "PlagSini_Platform_Overview_v1.0.md",
            out_path=ROOT / "PlagSini_Platform_Overview_v1.0.docx",
            title="PlagSini Platform Overview",
            subtitle="EV Charging Management Platform — Partner Introduction",
            date_str="1 July 2026",
        )
        build(
            src_path=ROOT / "PlagSini_OCPP_1.6J_Integration_Guide_v1.0.md",
            out_path=ROOT / "PlagSini_OCPP_1.6J_Integration_Guide_v1.0.docx",
            title="OCPP 1.6J Integration Guide",
            subtitle="Charger Hardware Vendor Onboarding — PlagSini EV Charging",
            date_str="1 July 2026",
        )
    else:
        build()
